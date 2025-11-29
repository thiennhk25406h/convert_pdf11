import streamlit as st
import fitz
from docx import Document
from docx.shared import Pt


def convert_pdf_to_docx(pdf_bytes):
    docx_file = "output.docx"
    pdf = fitz.open(stream=pdf_bytes, filetype="pdf")

    word = Document()
    style = word.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    for page in pdf:
        text = page.get_text()
        lines = text.split("\n")

        for line in lines:
            clean = line.strip()
            if not clean:
                continue

            # Câu hỏi
            if clean.lower().startswith("câu"):
                p = word.add_paragraph()
                run = p.add_run(clean)
                run.bold = True

            # Đáp án A/B/C/D
            elif clean[:1].upper() in ["A", "B", "C", "D"] and clean[1:2] == ".":
                word.add_paragraph("\t" + clean)

            else:
                word.add_paragraph(clean)

    word.save(docx_file)
    return docx_file


# GIAO DIỆN WEB
st.title("Chuyển PDF → DOCX (định dạng trắc nghiệm)")
uploaded = st.file_uploader("Tải lên file PDF", type=["pdf"])

if uploaded:
    st.success("File đã tải lên, bắt đầu xử lý…")

    docx_path = convert_pdf_to_docx(uploaded.read())

    with open(docx_path, "rb") as f:
        st.download_button("Tải file DOCX", f, file_name="converted.docx")
